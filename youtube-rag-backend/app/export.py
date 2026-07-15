from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path

import requests
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

from app.rate_limiter import rate_limit
from app.media_manager import get_media_meta

log = logging.getLogger("export")
router = APIRouter()

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ExportSection(BaseModel):
    heading: str
    bullets: list[str]

class ExportPayload(BaseModel):
    video_id:      str
    video_title:   str
    channel_name:  str = ""
    video_url:     str = ""
    summary:       str = ""
    sections:      list[ExportSection] = []
    key_takeaways: list[str] = []
    timestamps:    list[dict] = []


def generate_docx_from_payload(payload: dict, output_path: str):
    doc = Document()

    title_para = doc.add_heading(payload["video_title"], 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    if payload.get("channel_name"):
        meta.add_run(f'{payload["channel_name"]}\n').bold = True
    if payload.get("video_url"):
        meta.add_run(f'{payload["video_url"]}\n')
    meta.add_run(f'Generated at: {payload.get("generated_at", datetime.now().strftime("%Y-%m-%d %H:%M"))}')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    if payload.get("summary"):
        doc.add_heading("Executive Summary", level=1)
        p = doc.add_paragraph(payload["summary"])
        p.paragraph_format.space_after = Pt(10)

    for section in payload.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        for bullet in section["bullets"]:
            p   = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet)
            run.font.size = Pt(11)

    if payload.get("key_takeaways"):
        doc.add_heading("Key Takeaways", level=1)
        for t in payload["key_takeaways"]:
            p      = doc.add_paragraph(style="List Bullet")
            run    = p.add_run(t)
            run.bold = True

    if payload.get("timestamps"):
        doc.add_heading("Important Moments", level=1)
        for ts in payload["timestamps"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f'{ts["display"]} ').bold = True
            p.add_run(f'— {ts.get("label", "")}')

    doc.save(output_path)


def _get_video_title(video_id: str) -> dict:
    try:
        res = requests.get(
            "https://www.youtube.com/oembed",
            params={"url": f"https://www.youtube.com/watch?v={video_id}", "format": "json"},
            timeout=5,
        )
        res.raise_for_status()
        data = res.json()
        return {"title": data.get("title", video_id), "channel": data.get("author_name", "")}
    except Exception:
        meta = get_media_meta(video_id)
        if meta:
            return {"title": meta.get("title", video_id), "channel": ""}
        return {"title": video_id, "channel": ""}


# Export endpoint — accepts full payload from browser

@router.post("/export/docx", dependencies=[Depends(rate_limit("export"))])
def export_docx(payload: ExportPayload):
    export_dir = Path("exports")
    export_dir.mkdir(exist_ok=True)
    out_path = export_dir / f"{payload.video_id}_notes.docx"

    # Fill in title/channel if browser didn't provide them
    title   = payload.video_title or _get_video_title(payload.video_id)["title"]
    channel = payload.channel_name or _get_video_title(payload.video_id).get("channel", "")

    docx_payload = {
        "video_title":   title,
        "channel_name":  channel,
        "video_url":     payload.video_url,
        "generated_at":  datetime.now().strftime("%Y-%m-%d %H:%M"),
        "summary":       payload.summary,
        "sections":      [{"heading": s.heading, "bullets": s.bullets} for s in payload.sections],
        "key_takeaways": payload.key_takeaways[:4] or ["See the content for key insights."],
        "timestamps":    payload.timestamps,
    }

    try:
        generate_docx_from_payload(docx_payload, str(out_path))
    except Exception as exc:
        log.exception("DOCX generation failed")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")

    safe_title = re.sub(r"[^\w\s-]", "", title)[:60].strip()
    filename   = f"{safe_title} — Study Notes.docx" if safe_title else "Study Notes.docx"

    return FileResponse(
        str(out_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
    