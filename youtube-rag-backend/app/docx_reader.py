from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument          
from langchain_core.documents import Document

MAX_WORDS: int = 20_000      
MIN_PARA_CHARS: int = 20     

def _clean(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", "", text)
    return re.sub(r"\s+", " ", text).strip()


def _word_count(text: str) -> int:
    return len(text.split())

class WordLimitExceeded(Exception):
    """Raised when the document exceeds MAX_WORDS before truncation."""


def load_docx_docs(
    file_path: str | Path,
    max_words: int = MAX_WORDS,
    truncate: bool = True,         
) -> tuple[list[Document], dict]:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    docx = DocxDocument(str(path))
    raw_paragraphs: list[str] = []
    for para in docx.paragraphs:
        text = _clean(para.text)
        if len(text) >= MIN_PARA_CHARS:
            raw_paragraphs.append(text)
    for table in docx.tables:
        for row in table.rows:
            cells = [_clean(c.text) for c in row.cells if _clean(c.text)]
            if cells:
                raw_paragraphs.append("  |  ".join(cells))
    docs: list[Document] = []
    total_words = 0
    truncated = False

    for i, text in enumerate(raw_paragraphs):
        wc = _word_count(text)

        if total_words + wc > max_words:
            if not truncate:
                raise WordLimitExceeded(
                    f"Document exceeds the {max_words:,}-word limit "
                    f"(reached at paragraph {i + 1})."
                )
            # Partial paragraph: include as many words as budget allows
            remaining = max_words - total_words
            if remaining > 10:
                text = " ".join(text.split()[:remaining])
                docs.append(Document(
                    page_content=text,
                    metadata={"start": i * 10, "source": "docx", "paragraph": i},
                ))
            truncated = True
            break

        docs.append(Document(
            page_content=text,
            metadata={
                "start": i * 10,
                "source": "docx",
                "paragraph": i,
                "page": i // 5  
            },
        ))
        total_words += wc

    meta = {
        "word_count":      total_words,
        "paragraph_count": len(docs),
        "truncated":       truncated,
        "max_words":       max_words,
    }

    return docs, meta